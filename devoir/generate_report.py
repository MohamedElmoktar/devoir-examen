#!/usr/bin/env python3
"""Script pour générer le rapport du projet en format DOCX."""

from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from datetime import datetime
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from matplotlib.patches import FancyBboxPatch, FancyArrowPatch
import os


def create_architecture_diagram():
    """Créer le schéma d'architecture du système."""
    fig, ax = plt.subplots(1, 1, figsize=(14, 10))
    ax.set_xlim(0, 14)
    ax.set_ylim(0, 10)
    ax.axis('off')

    # Couleurs
    cloud_color = '#E3F2FD'
    fog_color = '#FFF3E0'
    device_color = '#E8F5E9'
    kafka_color = '#FFCDD2'
    arrow_color = '#455A64'

    # ===== COUCHE CLOUD (haut) =====
    cloud_rect = FancyBboxPatch((0.5, 7.5), 13, 2, boxstyle="round,pad=0.05",
                                  facecolor=cloud_color, edgecolor='#1976D2', linewidth=2)
    ax.add_patch(cloud_rect)
    ax.text(7, 9.2, 'COUCHE CLOUD', fontsize=14, fontweight='bold', ha='center', color='#1976D2')

    # Cloud Aggregator
    agg_rect = FancyBboxPatch((5.5, 7.8), 3, 1.2, boxstyle="round,pad=0.02",
                               facecolor='white', edgecolor='#1976D2', linewidth=1.5)
    ax.add_patch(agg_rect)
    ax.text(7, 8.6, 'Cloud Aggregator', fontsize=10, fontweight='bold', ha='center')
    ax.text(7, 8.2, '(FedAvg)', fontsize=9, ha='center', style='italic')

    # Dashboard
    dash_rect = FancyBboxPatch((10, 7.8), 2.5, 1.2, boxstyle="round,pad=0.02",
                                facecolor='white', edgecolor='#1976D2', linewidth=1.5)
    ax.add_patch(dash_rect)
    ax.text(11.25, 8.5, 'Dashboard', fontsize=10, fontweight='bold', ha='center')
    ax.text(11.25, 8.1, 'Monitoring', fontsize=9, ha='center', style='italic')

    # ===== KAFKA (milieu) =====
    kafka_rect = FancyBboxPatch((2, 5), 10, 1.8),
    kafka_rect = FancyBboxPatch((2, 5), 10, 1.8, boxstyle="round,pad=0.05",
                                 facecolor=kafka_color, edgecolor='#C62828', linewidth=2)
    ax.add_patch(kafka_rect)
    ax.text(7, 6.5, 'Apache Kafka', fontsize=12, fontweight='bold', ha='center', color='#C62828')

    # Topics Kafka
    topics = [
        ('sensor-data\n-node-1', 2.5),
        ('sensor-data\n-node-2', 5),
        ('model-\nweights', 7.5),
        ('global-\nmodel', 10)
    ]
    for topic, x in topics:
        topic_rect = FancyBboxPatch((x, 5.2), 1.8, 1, boxstyle="round,pad=0.02",
                                     facecolor='white', edgecolor='#C62828', linewidth=1)
        ax.add_patch(topic_rect)
        ax.text(x + 0.9, 5.7, topic, fontsize=7, ha='center', va='center')

    # ===== COUCHE FOG (milieu-bas) =====
    fog_rect = FancyBboxPatch((0.5, 2.2), 13, 2.2, boxstyle="round,pad=0.05",
                               facecolor=fog_color, edgecolor='#E65100', linewidth=2)
    ax.add_patch(fog_rect)
    ax.text(7, 4.1, 'COUCHE FOG / EDGE', fontsize=14, fontweight='bold', ha='center', color='#E65100')

    # Fog Nodes
    for i, x in enumerate([3, 9], 1):
        fog_node = FancyBboxPatch((x, 2.5), 2.5, 1.4, boxstyle="round,pad=0.02",
                                   facecolor='white', edgecolor='#E65100', linewidth=1.5)
        ax.add_patch(fog_node)
        ax.text(x + 1.25, 3.5, f'Fog Node {i}', fontsize=10, fontweight='bold', ha='center')
        ax.text(x + 1.25, 3.1, 'Spark + SGD', fontsize=8, ha='center', style='italic')
        ax.text(x + 1.25, 2.7, 'Local Training', fontsize=8, ha='center', style='italic')

    # ===== COUCHE DEVICE (bas) =====
    device_rect = FancyBboxPatch((0.5, 0.3), 13, 1.5, boxstyle="round,pad=0.05",
                                  facecolor=device_color, edgecolor='#2E7D32', linewidth=2)
    ax.add_patch(device_rect)
    ax.text(7, 1.5, 'COUCHE DISPOSITIF (IoT)', fontsize=14, fontweight='bold', ha='center', color='#2E7D32')

    # Sensors
    for i, x in enumerate([2, 4.5, 7.5, 10], 1):
        sensor = FancyBboxPatch((x, 0.5), 1.5, 0.8, boxstyle="round,pad=0.02",
                                 facecolor='white', edgecolor='#2E7D32', linewidth=1)
        ax.add_patch(sensor)
        ax.text(x + 0.75, 0.9, f'Capteur {i}', fontsize=8, ha='center')

    # ===== FLÈCHES =====
    arrow_style = "Simple, tail_width=0.5, head_width=4, head_length=8"
    kw = dict(arrowstyle=arrow_style, color=arrow_color)

    # Capteurs vers Kafka
    ax.annotate("", xy=(3.4, 5), xytext=(3, 1.8), arrowprops=dict(arrowstyle='->', color='#2E7D32', lw=1.5))
    ax.annotate("", xy=(5.9, 5), xytext=(5.5, 1.8), arrowprops=dict(arrowstyle='->', color='#2E7D32', lw=1.5))

    # Kafka vers Fog Nodes (données capteurs)
    ax.annotate("", xy=(4, 4), xytext=(3.4, 5), arrowprops=dict(arrowstyle='->', color='#C62828', lw=1.5))
    ax.annotate("", xy=(10, 4), xytext=(5.9, 5), arrowprops=dict(arrowstyle='->', color='#C62828', lw=1.5))

    # Fog Nodes vers Kafka (model weights)
    ax.annotate("", xy=(8.4, 5.2), xytext=(5, 4), arrowprops=dict(arrowstyle='->', color='#E65100', lw=1.5))
    ax.annotate("", xy=(8.4, 5.2), xytext=(9.5, 4), arrowprops=dict(arrowstyle='->', color='#E65100', lw=1.5))

    # Kafka vers Cloud Aggregator (model weights)
    ax.annotate("", xy=(7, 7.8), xytext=(8.4, 6.2), arrowprops=dict(arrowstyle='->', color='#C62828', lw=1.5))

    # Cloud Aggregator vers Kafka (global model)
    ax.annotate("", xy=(10.9, 6.2), xytext=(8, 7.8), arrowprops=dict(arrowstyle='->', color='#1976D2', lw=1.5))

    # Kafka vers Fog Nodes (global model)
    ax.annotate("", xy=(4.5, 4), xytext=(10.9, 5.2), arrowprops=dict(arrowstyle='->', color='#1976D2', lw=1.5, connectionstyle="arc3,rad=-0.2"))
    ax.annotate("", xy=(10, 4), xytext=(10.9, 5.2), arrowprops=dict(arrowstyle='->', color='#1976D2', lw=1.5))

    # Dashboard connexion
    ax.annotate("", xy=(10, 8.4), xytext=(8.5, 8.4), arrowprops=dict(arrowstyle='->', color='#1976D2', lw=1.5))

    # ===== LÉGENDE =====
    legend_elements = [
        mpatches.Patch(facecolor=device_color, edgecolor='#2E7D32', label='Données Capteurs'),
        mpatches.Patch(facecolor=fog_color, edgecolor='#E65100', label='Poids Locaux'),
        mpatches.Patch(facecolor=cloud_color, edgecolor='#1976D2', label='Modèle Global'),
        mpatches.Patch(facecolor=kafka_color, edgecolor='#C62828', label='Apache Kafka'),
    ]
    ax.legend(handles=legend_elements, loc='upper left', fontsize=9)

    # Titre
    ax.text(7, 9.8, 'Architecture du Système d\'Apprentissage Fédéré',
            fontsize=16, fontweight='bold', ha='center')

    plt.tight_layout()

    # Sauvegarder l'image
    img_path = '/Volumes/external/federated-learning-kafka-spark/architecture_diagram.png'
    plt.savefig(img_path, dpi=150, bbox_inches='tight', facecolor='white')
    plt.close()

    return img_path

def create_report():
    doc = Document()

    # Configuration des styles
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # ===== PAGE DE TITRE =====
    title = doc.add_heading('Rapport Technique', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('Système d\'Apprentissage Fédéré avec Apache Kafka et Apache Spark')
    run.bold = True
    run.font.size = Pt(16)

    doc.add_paragraph()

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.add_run(f'Date: {datetime.now().strftime("%d/%m/%Y")}\n')
    info.add_run('Projet: federated-learning-kafka-spark')

    doc.add_page_break()

    # ===== TABLE DES MATIÈRES =====
    doc.add_heading('Table des Matières', level=1)
    toc_items = [
        '1. Résumé Exécutif',
        '2. Architecture du Système',
        '3. Composants Principaux',
        '4. Technologies Utilisées',
        '5. Fonctionnalités Implémentées',
        '6. Flux de Données',
        '7. Configuration',
        '8. Guide de Démarrage',
        '9. Structure des Fichiers',
        '10. Conclusion'
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.left_indent = Cm(1)

    doc.add_page_break()

    # ===== 1. RÉSUMÉ EXÉCUTIF =====
    doc.add_heading('1. Résumé Exécutif', level=1)
    doc.add_paragraph(
        'Ce projet implémente un système d\'apprentissage fédéré (Federated Learning) '
        'pour la détection d\'anomalies distribuée. Le système utilise Apache Kafka comme '
        'broker de messages distribué et Apache Spark pour le traitement distribué en streaming.'
    )
    doc.add_paragraph(
        'L\'algorithme de Federated Averaging (FedAvg) est utilisé pour entraîner un modèle '
        'global de machine learning à travers plusieurs nœuds de périphérie (edge nodes) '
        'sans centraliser les données brutes, préservant ainsi la confidentialité des données.'
    )

    # ===== 2. ARCHITECTURE DU SYSTÈME =====
    doc.add_heading('2. Architecture du Système', level=1)
    doc.add_paragraph(
        'Le système adopte une architecture hiérarchique à trois niveaux:'
    )

    # Ajouter le schéma d'architecture
    img_path = create_architecture_diagram()
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(img_path, width=Inches(6))
    doc.add_paragraph()
    caption = doc.add_paragraph('Figure 1: Schéma d\'architecture du système d\'apprentissage fédéré')
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.runs[0].italic = True
    doc.add_paragraph()

    # Niveau Device
    doc.add_heading('2.1 Couche Dispositif (Device Layer)', level=2)
    doc.add_paragraph(
        'Simulateurs de capteurs IoT générant des données de télémétrie. '
        'Ces capteurs simulent des équipements industriels produisant des mesures de vibration, '
        'température et pression.'
    )

    # Niveau Edge/Fog
    doc.add_heading('2.2 Couche Brouillard (Fog Layer)', level=2)
    doc.add_paragraph(
        'Nœuds de calcul distribués effectuant l\'entraînement local des modèles. '
        'Chaque nœud fog traite les données localement et envoie uniquement les poids '
        'du modèle au cloud, sans exposer les données brutes.'
    )

    # Niveau Cloud
    doc.add_heading('2.3 Couche Cloud (Cloud Layer)', level=2)
    doc.add_paragraph(
        'Agrégateur central combinant les poids des modèles provenant de tous les nœuds '
        'pour créer un modèle global unifié qui est ensuite redistribué aux nœuds.'
    )

    # ===== 3. COMPOSANTS PRINCIPAUX =====
    doc.add_heading('3. Composants Principaux', level=1)

    doc.add_heading('3.1 Producteur de Capteurs (Sensor Producer)', level=2)
    doc.add_paragraph('Fichier: producers/sensor_producer.py')
    bullets = [
        'Simule des capteurs IoT industriels pour les nœuds 1 et 2',
        'Génère des lectures synthétiques avec 5 caractéristiques: vibration (x, y, z), température, pression',
        'Implémente un taux d\'injection d\'anomalies de 10%',
        'Produit des messages JSON vers les topics Kafka: sensor-data-node-1, sensor-data-node-2'
    ]
    for bullet in bullets:
        doc.add_paragraph(bullet, style='List Bullet')

    doc.add_heading('3.2 Nœuds Fog (Fog Nodes)', level=2)
    doc.add_paragraph('Fichiers: fog_nodes/fog_node.py et fog_node_simple.py')
    bullets = [
        'Version Spark Streaming: utilise PySpark Structured Streaming pour un traitement évolutif',
        'Version simplifiée: Python pur avec consommateurs Kafka directs pour les tests',
        'Consomme les données des capteurs depuis Kafka',
        'Implémente l\'entraînement local par descente de gradient stochastique (SGD)',
        'Entraîne un modèle de régression logistique pour la classification binaire d\'anomalies',
        'Publie périodiquement les poids du modèle vers le topic model-weights',
        'S\'abonne aux mises à jour du modèle global depuis le topic global-model'
    ]
    for bullet in bullets:
        doc.add_paragraph(bullet, style='List Bullet')

    doc.add_heading('3.3 Agrégateur Cloud (Cloud Aggregator)', level=2)
    doc.add_paragraph('Fichiers: cloud_aggregator/aggregator.py et aggregator_simple.py')
    bullets = [
        'Implémente l\'algorithme Federated Averaging (FedAvg)',
        'Consomme les poids des modèles depuis le topic model-weights',
        'Agrège les poids de tous les nœuds participants avec pondération par nombre d\'échantillons',
        'Publie le modèle global vers le topic global-model',
        'Journalise les rounds d\'agrégation dans un fichier JSONL'
    ]
    for bullet in bullets:
        doc.add_paragraph(bullet, style='List Bullet')

    doc.add_heading('3.4 Tableau de Bord de Surveillance (Dashboard)', level=2)
    doc.add_paragraph('Fichier: monitor/dashboard.py')
    bullets = [
        'Deux modes: Live (streaming Kafka) et Offline (analyse de fichiers de log)',
        'Visualisation console montrant les graphiques ASCII de convergence de perte',
        'Statistiques par nœud (échantillons contribués, pertes)',
        'Détails des rounds d\'agrégation et tendances de perte en temps réel'
    ]
    for bullet in bullets:
        doc.add_paragraph(bullet, style='List Bullet')

    # ===== 4. TECHNOLOGIES UTILISÉES =====
    doc.add_heading('4. Technologies Utilisées', level=1)

    table = doc.add_table(rows=7, cols=3)
    table.style = 'Table Grid'

    # En-têtes
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Technologie'
    hdr_cells[1].text = 'Objectif'
    hdr_cells[2].text = 'Version'

    # Données
    data = [
        ('Apache Kafka', 'Broker de messages distribué', '7.4.0 (Confluent)'),
        ('Apache Spark', 'Traitement streaming distribué', '3.5.0 (PySpark)'),
        ('Zookeeper', 'Coordination Kafka', '7.4.0'),
        ('Python', 'Langage principal', '3.x'),
        ('NumPy', 'Calculs numériques', '1.26.2'),
        ('Docker Compose', 'Orchestration infrastructure', 'Intégré'),
    ]

    for i, (tech, obj, ver) in enumerate(data, 1):
        row_cells = table.rows[i].cells
        row_cells[0].text = tech
        row_cells[1].text = obj
        row_cells[2].text = ver

    doc.add_paragraph()
    doc.add_heading('Topics Kafka (4 au total):', level=2)
    topics = [
        'sensor-data-node-1: Lectures des capteurs du dispositif 1',
        'sensor-data-node-2: Lectures des capteurs du dispositif 2',
        'model-weights: Poids locaux des modèles provenant des nœuds fog',
        'global-model: Modèle global agrégé provenant du cloud'
    ]
    for topic in topics:
        doc.add_paragraph(topic, style='List Bullet')

    # ===== 5. FONCTIONNALITÉS IMPLÉMENTÉES =====
    doc.add_heading('5. Fonctionnalités Implémentées', level=1)

    doc.add_heading('5.1 Apprentissage Fédéré (Core)', level=2)
    features = [
        'Algorithme FedAvg avec moyenne pondérée par nombre d\'échantillons',
        'Collecte asynchrone des poids (pas de synchronisation stricte entre nœuds)',
        'Entraînement itératif multi-rounds'
    ]
    for f in features:
        doc.add_paragraph('✓ ' + f, style='List Bullet')

    doc.add_heading('5.2 Tolérance aux Pannes', level=2)
    features = [
        'Participation partielle des nœuds (continue avec un sous-ensemble)',
        'Nombre minimum configurable de nœuds pour l\'agrégation (MIN_NODES_FOR_AGGREGATION = 1)',
        'Récupération automatique lors du retour des nœuds'
    ]
    for f in features:
        doc.add_paragraph('✓ ' + f, style='List Bullet')

    doc.add_heading('5.3 Entraînement du Modèle', level=2)
    features = [
        'Régression logistique pour la détection binaire d\'anomalies',
        'Entraînement local basé sur SGD avec taux d\'apprentissage configurable (0.01)',
        'Normalisation des caractéristiques',
        'Traitement par lots via micro-batches',
        'Calcul de perte par entropie croisée binaire'
    ]
    for f in features:
        doc.add_paragraph('✓ ' + f, style='List Bullet')

    doc.add_heading('5.4 Surveillance et Observabilité', level=2)
    features = [
        'Tableau de bord en temps réel avec streaming Kafka',
        'Analyse hors ligne des logs',
        'Journalisation round par round en JSONL',
        'Suivi des contributions par nœud',
        'Historique des pertes'
    ]
    for f in features:
        doc.add_paragraph('✓ ' + f, style='List Bullet')

    # ===== 6. FLUX DE DONNÉES =====
    doc.add_heading('6. Flux de Données', level=1)

    doc.add_heading('6.1 Flux des Données des Capteurs', level=2)
    doc.add_paragraph(
        'Dispositif → Producteur → Kafka (sensor-data-node-X) → Nœud Fog → Entraînement du Modèle'
    )

    doc.add_heading('6.2 Flux des Poids', level=2)
    doc.add_paragraph(
        'Nœud Fog → Kafka (model-weights) → Agrégateur Cloud → FedAvg → '
        'Kafka (global-model) → Nœuds Fog'
    )

    doc.add_heading('6.3 Flux de Surveillance', level=2)
    doc.add_paragraph(
        'Kafka (global-model) → Dashboard → Visualisation console\n'
        'Agrégateur → logs/aggregation_log.jsonl → Dashboard hors ligne'
    )

    # ===== 7. CONFIGURATION =====
    doc.add_heading('7. Configuration', level=1)
    doc.add_paragraph('Tous les paramètres sont centralisés dans config/settings.py:')

    config_text = '''# Kafka
KAFKA_BOOTSTRAP_SERVERS = "localhost:29092"

# Entraînement
LEARNING_RATE = 0.01
BATCH_SIZE = 32
LOCAL_EPOCHS = 1
NUM_FEATURES = 5

# Apprentissage Fédéré
AGGREGATION_INTERVAL_SECONDS = 30
MIN_NODES_FOR_AGGREGATION = 1
WEIGHT_PUBLISH_INTERVAL_SECONDS = 10

# Modèle
ANOMALY_PROBABILITY = 0.1
NUM_FOG_NODES = 2

# Spark
SPARK_MASTER = "local[*]"'''

    p = doc.add_paragraph()
    run = p.add_run(config_text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)

    # ===== 8. GUIDE DE DÉMARRAGE =====
    doc.add_heading('8. Guide de Démarrage Rapide', level=1)

    steps = [
        ('Installer les dépendances', './run.sh install'),
        ('Démarrer l\'infrastructure', './run.sh start-infra'),
        ('Démarrer l\'agrégateur cloud', './run.sh aggregator'),
        ('Démarrer le nœud fog 1', './run.sh fog-node 1'),
        ('Démarrer le nœud fog 2', './run.sh fog-node 2'),
        ('Démarrer le producteur 1', './run.sh producer 1'),
        ('Démarrer le producteur 2', './run.sh producer 2'),
        ('Démarrer le dashboard', './run.sh dashboard'),
    ]

    for i, (desc, cmd) in enumerate(steps, 1):
        p = doc.add_paragraph()
        p.add_run(f'{i}. {desc}: ').bold = True
        run = p.add_run(cmd)
        run.font.name = 'Courier New'

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Alternative (mode démo): ').bold = True
    run = p.add_run('./run.sh demo')
    run.font.name = 'Courier New'
    doc.add_paragraph('Cette commande démarre tous les composants automatiquement avec tmux.')

    # ===== 9. STRUCTURE DES FICHIERS =====
    doc.add_heading('9. Structure des Fichiers', level=1)

    structure = '''federated-learning-kafka-spark/
├── config/
│   ├── __init__.py
│   └── settings.py              # Paramètres de configuration
├── producers/
│   ├── __init__.py
│   └── sensor_producer.py       # Simulateur de capteurs IoT
├── fog_nodes/
│   ├── __init__.py
│   ├── fog_node.py              # Version Spark Streaming
│   └── fog_node_simple.py       # Version Python pure
├── cloud_aggregator/
│   ├── __init__.py
│   ├── aggregator.py            # Version Spark Streaming
│   └── aggregator_simple.py     # Version Python pure
├── monitor/
│   ├── __init__.py
│   └── dashboard.py             # Surveillance live/offline
├── logs/
│   └── aggregation_log.jsonl    # Historique d'entraînement
├── docker-compose.yml           # Configuration Kafka/Zookeeper/Spark
├── requirements.txt             # Dépendances Python
├── run.sh                       # Script de gestion
├── README.md                    # Documentation
├── test_simple.py               # Test d'intégration
└── test_integration.py          # Test d'intégration complet'''

    p = doc.add_paragraph()
    run = p.add_run(structure)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)

    # ===== 10. CONCLUSION =====
    doc.add_heading('10. Conclusion', level=1)
    doc.add_paragraph(
        'Ce projet représente une implémentation complète et fonctionnelle d\'un système '
        'd\'apprentissage fédéré pour la détection d\'anomalies en environnement distribué. '
        'Les principales réalisations incluent:'
    )

    achievements = [
        'Architecture distribuée robuste avec Kafka et Spark',
        'Implémentation complète de l\'algorithme FedAvg',
        'Double implémentation (Spark Streaming et Python pur) pour flexibilité',
        'Système de surveillance en temps réel',
        'Tolérance aux pannes et récupération automatique',
        'Infrastructure conteneurisée avec Docker Compose',
        'Documentation et scripts de démarrage complets'
    ]
    for a in achievements:
        doc.add_paragraph(a, style='List Bullet')

    doc.add_paragraph()
    doc.add_paragraph(
        'Le système est prêt pour la production et démontre les meilleures pratiques '
        'en matière de conception de systèmes distribués, de tolérance aux pannes '
        'et de surveillance des applications.'
    )

    # Sauvegarder le document
    output_path = '/Volumes/external/federated-learning-kafka-spark/Rapport_Projet_Apprentissage_Federe.docx'
    doc.save(output_path)
    print(f'Rapport généré avec succès: {output_path}')
    return output_path

if __name__ == '__main__':
    create_report()
